# requirements_engineering/document_processing/parsers/docx_parser.py
from pathlib import Path
import docx
from .base_parser import BaseParser

class DocxParser(BaseParser):
    """Parser for Word documents"""
    
    def parse(self, file_path: Path) -> str:
        """Extract text from DOCX file"""
        doc = docx.Document(file_path)
        full_text = []
        
        # Extract text from paragraphs
        for para in doc.paragraphs:
            full_text.append(para.text)
            
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
                    
        return "\n".join(full_text)
    
    def get_supported_extensions(self) -> list[str]:
        return [".docx", ".doc"]
